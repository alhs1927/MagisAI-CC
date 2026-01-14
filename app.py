import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import re

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="Magis AI - Kanisius",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Inisialisasi Session State
if 'result_text' not in st.session_state:
    st.session_state.result_text = ""
if 'topic_context' not in st.session_state:
    st.session_state.topic_context = ""

# --- 2. KAMUS BAHASA & PROMPT (IGNATIUS + CANISIUS) ---
TRANS = {
    "ID": {
        "title_sub": "Mitra Diskresi Guru Ignasian",
        "sidebar_settings": "Pengaturan",
        "lbl_lang": "Bahasa",
        "lbl_theme": "Tampilan",
        "lbl_tone": "Gaya Bahasa AI",
        "opt_tone": ["Reflektif (Spirit Jesuit)", "Akademis (4C+1L)", "Pastoral (Cura Personalis)"],
        "lbl_key": "Kunci Akses Google",
        "lbl_model": "Model Kecerdasan (Ganti jika Error)",
        "lbl_menu": "Pilih Modul Formasi",
        "menu_opt": ["1. Konteks (Cura Personalis)", "2. Desain RPP (IPP & 4C+1L)", "3. Refleksi (Examen & Magis)", "4. Asisten Makalah (Ignatius & Kanisius)"],
        "btn_analyze": "Analisis Profil Siswa",
        "btn_rpp": "Desain Pembelajaran PPI",
        "btn_reflect": "Mulai Diskresi",
        "btn_makalah": "Susun Draf Makalah",
        "btn_dl_word": "📥 Unduh Dokumen (.docx) - Format Rapi",
        "loading": "✨ Sedang menimbang (Diskresi) dalam terang roh...",
        "empty_warning": "⚠️ Mohon isi data untuk memulai proses.",
        "key_warning": "🔒 Masukkan Google API Key di sidebar kiri untuk memulai.",
        
        # --- MODUL 1: CURA PERSONALIS ---
        "m1_t": "📘 Konteks (Cura Personalis)", 
        "m1_l1": "Profil Unik Siswa / Dinamika Kelas:", 
        "m1_p1": "Ceritakan karakter, latar belakang, atau tantangan siswa (akademis/perilaku)...", 
        "m1_l2": "Fokus Masalah / Situasi:", 
        "m1_sys": """PERAN: Pendidik Ignasian yang menerapkan 'Cura Personalis' (perhatian pribadi).
        TUGAS: Analisis profil siswa melalui lensa 4C + 1L.
        FORMAT:
        1. **Diagnosa Situasi:** Apa yang terjadi pada batin siswa? (Lihat secara utuh).
        2. **Strategi Pendampingan:** Pendekatan personal apa yang bisa dilakukan?
        3. **Integrasi Nilai:** Fokuskan pada Conscience & Compassion.
        4. **Spirit Persevera:** Bagaimana guru bisa sabar mendampingi?
        Gunakan Tabel untuk memetakan Masalah vs Solusi.""",
        
        # --- MODUL 2: IPP (PPI) ---
        "m2_t": "📙 Desain Pembelajaran (PPI)", 
        "m2_l1": "Topik / Materi Pembelajaran:", 
        "m2_l2": "Durasi & Target Formasi:", 
        "m2_sys": """PERAN: Perancang Kurikulum Ignasian (IPP).
        TUGAS: Buat RPP yang membentuk siswa menjadi Leader yang memiliki 4C.
        SIKLUS IPP:
        1. **Context:** Siapa pembelajarnya?
        2. **Experience:** Aktivitas (Competence & Compassion).
        3. **Reflection:** Mengasah Hati Nurani (Conscience). Pertanyaan makna.
        4. **Action:** Aksi nyata (Leadership).
        5. **Evaluation:** Evaluasi pemahaman & hati.
        Gunakan Tabel.""",
        
        # --- MODUL 3: EXAMEN ---
        "m3_t": "📗 Refleksi Batin (Examen)", 
        "m3_l1": "Peristiwa / Kegelisahan / Topik Refleksi:", 
        "m3_sys": """PERAN: Pembimbing Rohani dengan semangat St. Ignatius Loyola.
        TUGAS: Pandu Examen Conscientiae (Pemeriksaan Batin).
        LANGKAH:
        1. **Gratitude (Syukur):** Menyadari kehadiran Tuhan.
        2. **Grace (Mohon Terang):** Memohon kejernihan.
        3. **Review (Tinjauan):** Melihat kembali peristiwa.
        4. **Repent (Penyesalan):** Mohon ampun atas kekurangan kasih.
        5. **Resolve (Niat):** Langkah Magis ke depan.""",

        # --- MODUL 4: MAKALAH (IGNATIUS & KANISIUS) ---
        "m4_t": "📘 Asisten Makalah (Jesuit Scholar)", 
        "m4_l1": "Topik / Judul Makalah:", 
        "m4_l2": "Argumen Utama (Diskresi & Kompetensi):", 
        "m4_p2": "Contoh: Teknologi harus diarahkan untuk memuliakan Tuhan dan melayani sesama (AMDG)...",
        "m4_sys": """PERAN: Cendekiawan Jesuit yang memadukan visi Visioner St. Ignatius Loyola (Pendiri, Ahli Diskresi) dan Intelektualitas St. Petrus Kanisius (Doktor Gereja, Edukator).
        
        TUGAS: Buat makalah yang cerdas (Competence), bijaksana (Conscience), dan memuliakan Tuhan (AMDG).
        
        STRUKTUR & ISI:
        1. **Judul:** Mencerminkan kedalaman akademis dan rohani.
        2. **Pendahuluan:** Konteks masalah.
        3. **Pembahasan (Dialog Iman & Ilmu):**
           - Gunakan logika deduktif yang kuat (Ciri St. Kanisius).
           - Masukkan unsur **DISKRESI** (Penimbangan Roh): Mengapa argumen ini dipilih? Apakah ini membawa kebaikan lebih besar (Magis)? (Ciri St. Ignatius).
           - Terapkan prinsip **Tantum Quantum** (Sejauh mana hal ini berguna bagi tujuan akhir).
        4. **Kesimpulan:** Arahkan pada 'Ad Maiorem Dei Gloriam' (Demi Kemuliaan Allah yang Lebih Besar) dan pelayanan sesama.
        
        GAYA BAHASA: Naratif, mengalir, persuasif, namun tetap rendah hati.""",
    },
    "EN": {
        "title_sub": "Ignatian Pedagogical Partner",
        "sidebar_settings": "Settings",
        "lbl_lang": "Language",
        "lbl_theme": "Theme",
        "lbl_tone": "AI Tone & Style",
        "opt_tone": ["Reflective (Jesuit Spirit)", "Academic (4C+1L)", "Pastoral (Cura Personalis)"],
        "lbl_key": "Google Access Key",
        "lbl_model": "Intelligence Model (Change if Error)",
        "lbl_menu": "Select Formation Module",
        "menu_opt": ["1. Context (Cura Personalis)", "2. Lesson Design (IPP & 4C+1L)", "3. Reflection (Examen & Magis)", "4. Paper Assistant (Ignatius & Canisius)"],
        "btn_analyze": "Analyze Student Profile",
        "btn_rpp": "Design IPP Lesson",
        "btn_reflect": "Start Discernment",
        "btn_makalah": "Draft Paper",
        "btn_dl_word": "📥 Download Document (.docx)",
        "loading": "✨ Discerning in the spirit of Magis...",
        "empty_warning": "⚠️ Please provide input.",
        "key_warning": "🔒 Please enter API Key in sidebar to start.",
        
        "m1_t": "📘 Context (Cura Personalis)", 
        "m1_l1": "Student Profile:", 
        "m1_p1": "Describe character, background...", 
        "m1_l2": "Situation:", 
        "m1_sys": "ROLE: Canisian Educator. Analyze student via 4C+1L. Focus on Cura Personalis (Care for the whole person). Use Tables.",
        
        "m2_t": "📙 Lesson Design (IPP)", 
        "m2_l1": "Topic:", 
        "m2_l2": "Goal:", 
        "m2_sys": "ROLE: IPP Designer. Cycle: Context -> Experience -> Reflection -> Action -> Evaluation. Build 4C+1L. Use Tables.",
        
        "m3_t": "📗 Reflection (Examen)", 
        "m3_l1": "Topic:", 
        "m3_sys": "ROLE: Spiritual Director. Guide Examen (Loyola style): Gratitude, Grace, Review, Repent, Resolve.",

        "m4_t": "📘 Academic Paper Assistant", 
        "m4_l1": "Paper Topic:", 
        "m4_l2": "Main Thesis (Discernment & Competence):", 
        "m4_p2": "E.g., Technology serves AMDG...",
        "m4_sys": "ROLE: Jesuit Scholar combining St. Ignatius (Discernment, AMDG) and St. Canisius (Intellect, Clarity). Write a paper that is logically sound AND spiritually discerning. End with Ad Maiorem Dei Gloriam.",
    }
}

# --- 3. LOGIKA WORD PROCESSING (Arial 12, Justify) ---

def clean_text(text):
    text = text.replace('**', '').replace('__', '')
    text = text.replace('```', '')
    text = re.sub(r'\$(.*?)\$', r'\1', text) 
    return text

def process_markdown_to_docx(doc, text):
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    lines = text.split('\n')
    table_buffer = [] 
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')]
            if len(cells) > 2: cells = cells[1:-1] 
            if '---' in cells[0]: continue 
            table_buffer.append(cells)
        else:
            if table_buffer:
                rows = len(table_buffer)
                cols = len(table_buffer[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_buffer):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            cell_p = table.rows[i].cells[j].paragraphs[0]
                            run = cell_p.add_run(clean_text(cell_text))
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            cell_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
                doc.add_paragraph() 
                table_buffer = [] 

            if stripped:
                clean_line = clean_text(stripped)
                if stripped.startswith('### '):
                    h = doc.add_heading(clean_line.replace('### ', ''), level=2)
                    h.runs[0].font.name = 'Arial'
                    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif stripped.startswith('## '):
                    h = doc.add_heading(clean_line.replace('## ', ''), level=1)
                    h.runs[0].font.name = 'Arial'
                    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    p = doc.add_paragraph(clean_line.replace('- ', '').replace('* ', ''), style='List Bullet')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if p.runs: p.runs[0].font.name = 'Arial'
                elif re.match(r'^\d+\.', stripped):
                    p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', clean_line), style='List Number')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if p.runs: p.runs[0].font.name = 'Arial'
                else:
                    p = doc.add_paragraph(clean_line)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if p.runs: p.runs[0].font.name = 'Arial'

    if table_buffer:
        rows = len(table_buffer)
        cols = len(table_buffer[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_buffer):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell_p = table.rows[i].cells[j].paragraphs[0]
                    run = cell_p.add_run(clean_text(cell_text))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

def create_docx(content, topic, lang_key):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    header_text = 'MAGIS AI RESULT' if lang_key == 'EN' else 'MAGIS AI - HASIL DISKRESI'
    h = doc.add_heading(header_text, 0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    runner = p.add_run(f'Topic: {topic}')
    runner.bold = True
    runner.font.name = 'Arial'
    runner.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n')
    
    process_markdown_to_docx(doc, content)
    
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_text = 'Dicetak oleh Magis AI - Kolese Kanisius'
    f_run = footer_para.add_run(f'\n--- {footer_text} ---')
    f_run.font.name = 'Arial'
    f_run.font.size = Pt(9)
    f_run.italic = True
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    bio = BytesIO()
    doc.save(bio)
    return bio

def get_gemini_response(api_key, model_name, system_instruction, user_prompt, tone, lang):
    if not api_key: return None
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        lang_instruction = "Use standard Indonesian." if lang == "ID" else "Use professional English."
        
        # PROMPT UTAMA: 4C + 1L & STYLE IGNATIUS + KANISIUS
        format_instruction = """
        FORMAT: Use Markdown Tables for structured data. Use standard paragraphs (narrative flow).
        CORE VALUES (Integrate these): 
        - Competence (Academic excellence - Canisius)
        - Conscience (Moral discernment - Ignatius)
        - Compassion (Empathy for others)
        - Commitment (Persevera/Grit)
        - Leadership (Service/AMDG)
        - Magis (Striving for the better)
        - Finding God in All Things
        """
        
        full_sys = f"ROLE: {system_instruction}\nTONE: {tone}\nLANGUAGE: {lang_instruction}\n{format_instruction}"
        response = model.generate_content(f"{full_sys}\n\nTASK: {user_prompt}")
        text = response.text
        if text.startswith("```"): text = text.replace("```markdown", "").replace("```", "")
        return text
    except Exception as e: return f"Error: {str(e)}"

# --- 4. CSS (PERBAIKAN LIGHT MODE & UI) ---
def inject_custom_css(theme):
    if theme == "Gelap":
        vars = """
            --bg-color: #0E1117; 
            --sidebar-bg: #161B22; 
            --text-color: #E6EDF3; 
            --input-bg: #0d1117; 
            --input-border: #30363D; 
            --card-bg: #161B22; 
            --primary-color: #4285F4;
            --header-color: #FFFFFF;
        """
    else:
        # CLEAN LIGHT MODE (High Contrast)
        vars = """
            --bg-color: #FFFFFF; 
            --sidebar-bg: #F8F9FA; 
            --text-color: #000000; 
            --input-bg: #FFFFFF; 
            --input-border: #BDC3C7; 
            --card-bg: #F0F2F6; 
            --primary-color: #0047AB; 
            --header-color: #333333;
        """

    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;700;800&display=swap');
    
    :root {{ {vars} }}
    
    html, body, .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {{ 
        background-color: var(--bg-color) !important; 
        font-family: 'Inter', sans-serif; 
        color: var(--text-color) !important;
    }}
    
    h1, h2, h3, h4, h5, h6 {{ color: var(--header-color) !important; }}
    p, li, span, div, label {{ color: var(--text-color) !important; }}
    
    section[data-testid="stSidebar"] {{ 
        background-color: var(--sidebar-bg); 
        border-right: 1px solid var(--input-border); 
    }}
    section[data-testid="stSidebar"] h1, section[data-testid="stSidebar"] h2, section[data-testid="stSidebar"] h3 {{
        color: var(--text-color) !important;
    }}
    
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div {{ 
        background-color: var(--input-bg) !important; 
        color: var(--text-color) !important; 
        border: 1px solid var(--input-border) !important; 
        border-radius: 8px; 
    }}
    ::placeholder {{ color: var(--text-color); opacity: 0.7; }}
    
    div.stButton > button {{ 
        background: linear-gradient(90deg, #4285F4, #34A853); 
        color: white !important; 
        border: none; 
        border-radius: 50px; 
        font-weight: 600; 
        box-shadow: 0 2px 5px rgba(0,0,0,0.1); 
        transition: all 0.2s; 
    }}
    div.stButton > button:hover {{ transform: scale(1.02); box-shadow: 0 4px 8px rgba(0,0,0,0.2); }}
    
    .result-card {{ 
        background-color: var(--card-bg); 
        border: 1px solid var(--input-border); 
        border-radius: 12px; 
        padding: 40px; 
        box-shadow: 0 4px 12px rgba(0,0,0,0.05); 
        margin-top: 20px; 
        text-align: justify; 
        line-height: 1.8;
        color: var(--text-color) !important;
    }}
    
    .result-card h1, .result-card h2, .result-card h3 {{
        color: var(--primary-color) !important;
        margin-top: 25px;
        margin-bottom: 15px;
        border-bottom: 1px solid var(--input-border);
        padding-bottom: 5px;
        text-align: left;
    }}
    
    table {{ width: 100%; border-collapse: collapse; border: 1px solid var(--input-border); margin: 20px 0; }}
    th, td {{ padding: 12px; border: 1px solid var(--input-border); text-align: left; color: var(--text-color); }}
    th {{ background-color: rgba(66, 133, 244, 0.1); font-weight: bold; }}

    .main-title {{ 
        background: linear-gradient(90deg, #4285F4, #EA4335, #FBBC05, #34A853); 
        -webkit-background-clip: text; 
        -webkit-text-fill-color: transparent; 
        font-weight: 800; 
        font-size: 3rem; 
        text-align: center; 
        margin-bottom: 0.5rem;
    }}
    </style>
    """, unsafe_allow_html=True)

# --- 5. TAMPILAN UI ---
with st.sidebar:
    st.markdown("<div style='text-align:center; margin-bottom:20px;'><img src='https://i.imgur.com/UUCgyfV.png' width='90'></div>", unsafe_allow_html=True)
    lang_opt = st.radio("Bahasa", ["Indonesia 🇮🇩", "English 🇺🇸"], horizontal=True, label_visibility="collapsed")
    L_CODE = "ID" if "Indonesia" in lang_opt else "EN"
    TXT = TRANS[L_CODE] 
    
    st.markdown(f"### ⚙️ {TXT['sidebar_settings']}")
    theme_opt = st.radio(TXT["lbl_theme"], ["Modern Dark 🌑", "Clean Light ☀️"])
    THEME_VAL = "Gelap" if "Dark" in theme_opt else "Terang"
    
    tone_idx = st.selectbox(TXT["lbl_tone"], TXT["opt_tone"])
    st.divider()
    
    api_key = st.text_input(TXT["lbl_key"], type="password")
    models = []
    if api_key:
        try:
            genai.configure(api_key=api_key)
            models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        except: pass
    
    # Dropdown Model (Solusi Error 429)
    sel_model = st.selectbox(TXT["lbl_model"], models) if models else None
    
    st.divider()
    menu_sel = st.radio(TXT["lbl_menu"], TXT["menu_opt"])
    menu_idx = TXT["menu_opt"].index(menu_sel)
    
    # --- CREDIT TITLE ---
    st.markdown(f"""
    <div style='margin-top:3rem;text-align:center;font-size:0.7rem;opacity:0.7;line-height:1.5;color:var(--text-color);'>
    <strong>MAGIS AI v12.0 (JESUIT SCHOLAR)</strong><br>
    Design by: Albertus Henny Setyawan<br>
    Kolese Kanisius Jakarta | 2026
    </div>
    """, unsafe_allow_html=True)

inject_custom_css(THEME_VAL)

st.markdown(f"<div class='main-title'>Magis AI</div><div style='text-align:center;color:grey;margin-bottom:30px;font-style:italic;'>{TXT['title_sub']}</div>", unsafe_allow_html=True)

if not api_key:
    st.info(TXT["key_warning"])
    st.stop()

# --- INPUT AREA ---
with st.container():
    if menu_idx == 0: 
        st.markdown(f"<h3 style='color:#4285F4;'>{TXT['m1_t']}</h3>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        in_1 = c1.text_area(TXT['m1_l1'], placeholder=TXT['m1_p1'], height=150)
        in_2 = c2.text_input(TXT['m1_l2'])
        if st.button(TXT['btn_analyze']):
            if in_1:
                prompt = f"{TXT['m1_l1']} {in_1} | {TXT['m1_l2']} {in_2}"
                with st.spinner(TXT['loading']):
                    res = get_gemini_response(api_key, sel_model, TXT['m1_sys'], prompt, tone_idx, L_CODE)
                    st.session_state.result_text = res
                    st.session_state.topic_context = prompt
            else: st.warning(TXT['empty_warning'])

    elif menu_idx == 1: 
        st.markdown(f"<h3 style='color:#FBBC05;'>{TXT['m2_t']}</h3>", unsafe_allow_html=True)
        c1, c2 = st.columns([2, 1])
        in_1 = c1.text_input(TXT['m2_l1'])
        in_2 = c2.selectbox(TXT['m2_l2'], ["1 JP (45')", "2 JP (90')", "Block Project (Project Based)"])
        if st.button(TXT['btn_rpp']):
            if in_1:
                prompt = f"Topik Pembelajaran: {in_1} | Durasi: {in_2}"
                with st.spinner(TXT['loading']):
                    res = get_gemini_response(api_key, sel_model, TXT['m2_sys'], prompt, tone_idx, L_CODE)
                    st.session_state.result_text = res
                    st.session_state.topic_context = prompt
            else: st.warning(TXT['empty_warning'])

    elif menu_idx == 2:
        st.markdown(f"<h3 style='color:#34A853;'>{TXT['m3_t']}</h3>", unsafe_allow_html=True)
        in_1 = st.text_area(TXT['m3_l1'], height=100)
        if st.button(TXT['btn_reflect']):
            if in_1:
                prompt = f"Bahan Refleksi: {in_1}"
                with st.spinner(TXT['loading']):
                    res = get_gemini_response(api_key, sel_model, TXT['m3_sys'], prompt, tone_idx, L_CODE)
                    st.session_state.result_text = res
                    st.session_state.topic_context = prompt
            else: st.warning(TXT['empty_warning'])

    elif menu_idx == 3: # MODUL 4: MAKALAH (IGNATIUS & KANISIUS)
        st.markdown(f"<h3 style='color:#4285F4;'>{TXT['m4_t']}</h3>", unsafe_allow_html=True)
        in_1 = st.text_input(TXT['m4_l1'])
        in_2 = st.text_area(TXT['m4_l2'], placeholder=TXT['m4_p2'], height=150)
        
        len_opt = st.radio("Target Kedalaman Tulisan:", 
                           ["Ringkas (Poin 4C)", "Sedang (Standard Makalah)", "Mendalam (Analisis Diskresi Penuh)"], 
                           horizontal=True)
        
        if st.button(TXT['btn_makalah']):
            if in_1 and in_2:
                # Prompt Khusus: Gabungan St. Ignatius & St. Kanisius
                prompt = f"""
                TOPIK: {in_1}
                ARGUMEN UTAMA (TESIS): {in_2}
                TARGET KEDALAMAN: {len_opt}
                
                INSTRUKSI KHUSUS (JESUIT STYLE):
                1. **Competence (St. Kanisius):** Sajikan data dan argumen logis yang tak terbantahkan.
                2. **Conscience (St. Ignatius):** Lakukan DISKRESI dalam pembahasan. Timbang baik-buruk moralnya.
                3. **Finding God in All Things:** Ajak pembaca menemukan nilai Ilahi dalam topik ini.
                4. **AMDG:** Tutup dengan bagaimana hal ini memuliakan Tuhan (Ad Maiorem Dei Gloriam).
                5. Gaya bahasa: Naratif, mengalir, tanpa list berlebihan.
                """
                with st.spinner(TXT['loading']):
                    res = get_gemini_response(api_key, sel_model, TXT['m4_sys'], prompt, tone_idx, L_CODE)
                    st.session_state.result_text = res
                    st.session_state.topic_context = f"Makalah: {in_1}"
            else: st.warning(TXT['empty_warning'])

# --- OUTPUT AREA ---
if st.session_state.result_text:
    st.markdown("---")
    st.markdown(f"<div class='result-card'>{st.session_state.result_text}</div>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    
    docx_file = create_docx(st.session_state.result_text, st.session_state.topic_context, L_CODE)
    
    st.download_button(
        label=TXT['btn_dl_word'],
        data=docx_file.getvalue(),
        file_name=f"MagisAI_{L_CODE}_JesuitScholar.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )